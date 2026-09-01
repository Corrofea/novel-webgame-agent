#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从 characters.json 生成 Word 人物卡片册（每人一页）。

用法:
    python character_cards.py --input outputs/characters.json --output outputs/

依赖: pip install python-docx
书名/作者从 JSON 的 book/author 字段读取（不再硬编码）。
"""
import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT = '微软雅黑'
ACCENT = RGBColor(0x8B, 0x00, 0x00)
ACCENT_LIGHT = RGBColor(0x58, 0x8E, 0x32)
TEXT = RGBColor(0x2A, 0x33, 0x2E)
SUB = RGBColor(0x5A, 0x78, 0x64)


def _set_font(run, size, bold=False, color=TEXT):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def join_list(items):
    if not items:
        return ['未提及']
    return items if isinstance(items, list) else [str(items)]


def create_card(doc, char):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(char.get('name', '未命名')), 22, bold=True, color=ACCENT)
    p.paragraph_format.space_after = Pt(6)

    aliases = char.get('aliases') or []
    if aliases and aliases[0] != '无':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run('别名：' + ' / '.join(aliases)), 10, color=SUB)
        p.paragraph_format.space_after = Pt(12)
    else:
        doc.add_paragraph()

    fields = [
        ('身份', ' / '.join(filter(None, [char.get('gender', ''), char.get('identity', '')])) or '未提及'),
        ('性格特点', join_list(char.get('traits'))),
        ('主要经历', join_list(char.get('experiences'))),
        ('人物关系', join_list(char.get('relationships'))),
        ('结局/现状', char.get('ending') or '未提及'),
        ('备注', join_list(char.get('notes'))),
    ]
    for label, items in fields:
        p = doc.add_paragraph()
        _set_font(p.add_run(label), 12, bold=True, color=ACCENT_LIGHT)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _set_font(p.add_run('• ' + str(item)), 10)
            p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph()


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--input', required=True, help='characters.json 路径')
    ap.add_argument('--output', required=True, help='输出目录')
    args = ap.parse_args()

    data = json.loads(Path(args.input).read_text(encoding='utf-8'))
    book = data.get('book', '未命名')
    author = data.get('author', '')
    chars = data.get('characters', [])

    doc = Document()
    section = doc.sections[0]
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        try:
            setattr(section, attr, Pt(42))
        except Exception:
            pass

    # 封面
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(20)
    _set_font(p.add_run(f'《{book}》人物卡片'), 28, bold=True, color=ACCENT)
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(f'{author} 著'), 14, color=SUB)

    for char in chars:
        doc.add_page_break()
        create_card(doc, char)

    out_dir = Path(args.output)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f'{book}_人物卡片.docx'
    doc.save(str(out_path))
    print(f'已生成: {out_path}')


if __name__ == '__main__':
    main()
